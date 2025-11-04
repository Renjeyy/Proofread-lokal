import os
import io
import re
import json
import zipfile
import difflib
import fitz  # PyMuPDF
import docx
import pandas as pd
import google.generativeai as genai
from flask import Flask, request, jsonify, render_template, send_file, make_response
from dotenv import load_dotenv
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Inches

# Muat environment variables (GOOGLE_API_KEY) dari file .env
load_dotenv()

# --- Konfigurasi App & AI ---
app = Flask(__name__)

try:
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError("GOOGLE_API_KEY tidak ditemukan di file .env")
    
    genai.configure(api_key=api_key)
    # Ganti dengan model yang Anda inginkan
    model = genai.GenerativeModel('gemini-2.5-pro') 
except Exception as e:
    print(f"Error saat mengkonfigurasi Google AI: {e}")
    # Aplikasi akan tetap berjalan, tetapi endpoint AI akan gagal

def _extract_text_with_pages(file_bytes, file_extension):
    """Mengekstrak teks dari file PDF atau DOCX (versi backend)."""
    pages_content = []
    
    if file_extension == 'pdf':
        try:
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            for page_num, page in enumerate(pdf_document):
                pages_content.append({"halaman": page_num + 1, "teks": page.get_text()})
            pdf_document.close()
        except Exception as e:
            raise ValueError(f"Gagal membaca file PDF: {e}")
            
    elif file_extension == 'docx':
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            # Untuk docx, kita anggap sebagai 1 halaman besar
            pages_content.append({"halaman": 1, "teks": full_text})
        except Exception as e:
            raise ValueError(f"Gagal membaca file DOCX: {e}")
    else:
        raise ValueError("Format file tidak didukung. Harap unggah .pdf atau .docx")
        
    return pages_content

def _get_text_from_flask_file(file):
    """Utility untuk membaca file dari request Flask."""
    file_bytes = file.read()
    file_extension = file.filename.split('.')[-1].lower()
    return _extract_text_with_pages(file_bytes, file_extension)

def _get_full_text_from_file(file):
    """Utility untuk mendapatkan seluruh teks sebagai satu string."""
    pages = _get_text_from_flask_file(file)
    return "\n".join([page['teks'] for page in pages])


# --- Fungsi Logika AI (Sebagian besar disalin langsung) ---

def proofread_with_gemini(text_to_check):
    """Mengirim teks ke Gemini untuk proofreading (tanpa st)."""
    if not text_to_check or text_to_check.isspace():
        return []
    
    # NOTE: Prompt Anda SANGAT spesifik. Saya salin apa adanya.
    prompt = f"""
    Anda adalah seorang auditor dan ahli bahasa Indonesia yang sangat teliti. Anda diberikan dokumen dan tugas Anda adalah melakukan proofread pada teks berikut. Fokus pada:
    1. Memperbaiki kesalahan ketik (typo) agar semuanya sesuai dengan standar KBBI dan PUEBI.
    1. Kalau ada kata-kata yang tidak sesuai KBBI dan PUEBI, tolong jangan highlight semua kalimatnya, tapi cukup highlight kata-kata yang salah serta perbaiki kata-kata itu aja, jangan perbaiki semua kalimatnya
    3. Jika ada kata yang diitalic, biarkan saja
    4. Nama-nama yang diberi ini pastikan benar juga "Yullyan, I Made Suandi Putra, Laila Fajriani, Hari Sundoro, Bakhas Nasrani Diso, Rizky Ananda Putra, Wirawan Arief Nugroho, Lelya Novita Kusumawati, Ryani Ariesti Syafitri, Darmo Saputro Wibowo, Lucky Parwitasari, Handarudigdaya Jalanidhi Kuncaratrah, Fajar Setianto, Jaka Tirtana Hanafiah,  Muhammad Rosyid Ridho Muttaqien, Octovian Abrianto, Deny Sjahbani, Jihan Abigail, Winda Anggraini, Fadian Dwiantara, Aliya Anindhita Rachman"
    5. Fontnya arial dan jangan diganti. Khusus untuk judul paling atas, itu font sizenya 12 dan bodynya selalu 11
    6. Khusus "Indonesia Financial Group (IFG)", meskipun bahasa inggris, tidak perlu di italic
    7. Kalau ada kata yang sudah diberikan akronimnya di awal, maka di halaman berikut-berikutnya cukup akronimnya saja, tidak perlu ditulis lengkap lagi
    8. Pada bagian Nomor surat dan Penutup tidak perlu dicek, biarkan seperti itu
    9. Ketika Anda perbaiki, fontnya pastikan Arial dengan ukuran 11 juga (Tidak diganti)
    10. Pada kalimat "Indonesia Financial Group", jika terdapat kata typo "Finansial", tolong Anda sarankan untuk ganti ke "Financial"
    11. Yang benar adalah "Satuan Kerja Audit Internal", bukan "Satuan Pengendali Internal Audit"
    12. Jika terdapat kata "reviu", biarkan itu sebagai benar
    13. Kalau ada kata "IM", "ST", "SKAI", "IFG", "TV (Angka Romawi)", "RKAT", dan "RKAP", itu tidak perlu ditandai sebagai salah dan tidak perlu disarankan untuk italic / bold / underline
    14. Untuk nama modul seperti "Modul Sourcing, dll", itu tidak perlu italic
    15. Kalau ada kata dalam bahasa inggris yang masih masuk akal dan nyambung dengan kalimat yang dibahas, tidak perlu Anda sarankan untuk ganti ke bahasa indonesia
    16. Jika ada bahasa inggris dan akronimnya seperti "General Ledger (GL)", tolong dilakukan italic pada kata tersebut pada saat download file hasil revisinya, akronimnya tidak perlu diitalic
    17. Awal kalimat selalu dimulai dengan huruf kapital. Jika akhir poin diberi tanda ";", maka poin selanjutnya tidak perlu kapital
    18. Di file hasil revisi, Anda jangan ganti dari yang aslinya. Misalnya kalau ada kata yang diitalic di file asli, jangan Anda hilangkan italicnya
    19. Tolong perhatikan juga tanda bacanya, seperti koma, titik koma, titik, tanda hubung, dan lain-lain. Pastikan sesuai dan ada tanda titik di setiap akhir kalimat

    PENTING: Berikan hasil dalam format yang SANGAT KETAT. Untuk setiap kesalahan, gunakan format:
    [SALAH] kata atau frasa yang salah -> [BENAR] kata atau frasa perbaikan -> [KALIMAT] kalimat lengkap asli tempat kesalahan ditemukan

    Contoh:
    [SALAH] dikarenakan -> [BENAR] karena -> [KALIMAT] Hal itu terjadi dikarenakan kelalaian petugas.

    Jika tidak ada kesalahan sama sekali, kembalikan teks: "TIDAK ADA KESALAHAN"

    Berikut adalah teks yang harus Anda periksa:
    ---
    {text_to_check}
    """
    try:
        response = model.generate_content(prompt)
        pattern = re.compile(r"\[SALAH\]\s*(.*?)\s*->\s*\[BENAR\]\s*(.*?)\s*->\s*\[KALIMAT\]\s*(.*?)\s*(\n|$)", re.IGNORECASE | re.DOTALL)
        found_errors = pattern.findall(response.text)
        return [{"salah": salah.strip(), "benar": benar.strip(), "kalimat": kalimat.strip()} for salah, benar, kalimat, _ in found_errors]
    except Exception as e:
        print(f"Terjadi kesalahan saat menghubungi AI: {e}")
        # Kembalikan error sebagai list (agar JSON tidak error)
        return [{"salah": "ERROR", "benar": str(e), "kalimat": "Gagal menghubungi API"}]

def analyze_document_coherence(full_text):
    """Menganalisis koherensi (tanpa st)."""
    if not full_text or full_text.isspace():
        return []

    prompt = f"""
    Anda adalah seorang auditor ahli yang bertugas menganalisis struktur dan koherensi sebuah tulisan.
    Tugas Anda adalah membaca keseluruhan teks berikut dan mengidentifikasi setiap kalimat atau paragraf yang tidak koheren atau keluar dari topik utama di dalam sebuah sub-bagian.
    
    Untuk setiap ketidaksesuaian yang Anda temukan, lakukan hal berikut:
    1. Bacalah mengenai judul dari section atau subsection yang ada pada file tersebut
    2. Tentukan topik utama dari setiap section / subsection terutama isi paragrafnya.
    3. Identifikasi kalimat asli yang menyimpang dari topik tersebut dari yang telah Anda temukan pada section / subsection tersebut.
    4. Bila ada kalimat yang sekiranya memyimpang, Berikan saran dengan menghighlight kalimat tersebut untuk diulis ulang kalimat (rewording) tersebut agar relevan dan menyatu kembali dengan topik utamanya, sambil berusaha mempertahankan maksud aslinya jika memungkinkan.
    5. Kalau ada kata yang merupakan bahasa inggris, biarkan saja dan tidak perlu ditranslate ke bahasa indonesia, Anda cukup highlight kata tersebut
    6. Kalau ada kata yang tidak baku sesuai dengan standar KBBI, harap Anda perbaiki juga sehingga kata tersebut baku sesuai standar KBBI

    Berikan hasil dalam format yang SANGAT KETAT seperti di bawah ini. Ulangi format ini untuk setiap kalimat menyimpang yang Anda temukan:
    [TOPIK UTAMA] topik utama dari bagian tersebut -> [TEKS ASLI] kalimat asli yang tidak koheren -> [SARAN REVISI] versi kalimat yang sudah diperbaiki agar koheren

    Jika seluruh dokumen sudah koheren dan tidak ada masalah, kembalikan teks: "TIDAK ADA MASALAH KOHERENSI"

    Teks:
    ---
    {full_text}
    """
    try:
        response = model.generate_content(prompt)
        pattern = re.compile(r"\[TOPIK UTAMA\]\s*(.*?)\s*->\s*\[TEKS ASLI\]\s*(.*?)\s*->\s*\[SARAN REVISI\]\s*(.*?)\s*(\n|$)", re.IGNORECASE | re.DOTALL)
        found_issues = pattern.findall(response.text)
        return [{"topik": topik.strip(), "asli": asli.strip(), "saran": saran.strip()} for topik, asli, saran, _ in found_issues]
    except Exception as e:
        print(f"Terjadi kesalahan saat menghubungi AI: {e}")
        return [{"topik": "ERROR", "asli": str(e), "saran": "Gagal menghubungi API"}]

def get_structural_recommendations(full_text):
    """Menganalisis restrukturisasi (tanpa st)."""
    if not full_text or full_text.isspace():
        return []

    prompt = f"""
    Anda adalah seorang auditor ahli... (Salin prompt 'restrukturisasi' Anda ke sini) ...
    Berikan hasil dalam format JSON yang berisi sebuah list. Setiap objek harus memiliki tiga kunci: "misplaced_paragraph", "original_section", dan "recommended_section".

    Contoh Format JSON:
    [
      {{
        "misplaced_paragraph": "Selain itu, audit internal juga bertugas memeriksa laporan keuangan setiap kuartal...",
        "original_section": "Bab 2.1: Prosedur Whistleblowing",
        "recommended_section": "Bab 4.2: Peran Audit Internal"
      }}
    ]
    Jika dokumen sudah bagus, kembalikan list kosong: []

    Teks Dokumen:
    ---
    {full_text}
    """
    try:
        response = model.generate_content(prompt)
        cleaned_response = re.sub(r'```json\s*|\s*```', '', response.text.strip())
        return json.loads(cleaned_response)
    except Exception as e:
        print(f"Failed to Generate Response from AI: {e}")
        return [{"misplaced_paragraph": "ERROR", "original_section": str(e), "recommended_section": "Gagal menghubungi API"}]

# --- Fungsi Pemrosesan Dokumen (Disalin langsung) ---

def generate_revised_docx(file_bytes, errors):
    doc = docx.Document(io.BytesIO(file_bytes))
    for error in reversed(errors):
        salah = error["Kata/Frasa Salah"]
        benar = error["Perbaikan Sesuai KBBI"]
        for para in doc.paragraphs:
            if salah in para.text:
                # Logika penggantian teks sederhana
                # Catatan: Logika Anda sebelumnya untuk mempertahankan font rumit
                # dan mungkin perlu disempurnakan. Ini adalah versi sederhana.
                para.text = para.text.replace(salah, benar) 
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def generate_highlighted_docx(file_bytes, errors):
    doc = docx.Document(io.BytesIO(file_bytes))
    unique_salah = set(error["Kata/Frasa Salah"] for error in errors)
    for para in doc.paragraphs:
        for term in unique_salah:
            if term.lower() in para.text.lower():
                full_text = para.text
                para.clear()
                parts = re.split(f'({re.escape(term)})', full_text, flags=re.IGNORECASE)
                for part in parts:
                    if part:
                        run = para.add_run(part)
                        if part.lower() == term.lower():
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        # Pertahankan style dasar (jika perlu, logika bisa ditambahkan di sini)
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def create_zip_archive(revised_data, highlighted_data, original_filename):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr(f"revisi_{original_filename}", revised_data)
        zip_file.writestr(f"highlight_{original_filename}", highlighted_data)
    return zip_buffer.getvalue()

def extract_paragraphs(file_bytes):
    try:
        source_stream = io.BytesIO(file_bytes)
        doc = docx.Document(source_stream)
        return [p.text for p in doc.paragraphs if p.text.strip() != ""]
    except Exception as e:
        raise ValueError(f"Gagal membaca file docx: {e}")

def find_word_diff(original_para, revised_para):
    matcher = difflib.SequenceMatcher(None, original_para.split(), revised_para.split())
    diffs = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace' or tag == 'insert':
            diffs.append(" ".join(revised_para.split()[j1:j2]))
    return ", ".join(diffs) if diffs else "Perubahan Minor"

def create_comparison_docx(df):
    doc = Document()
    doc.add_heading('Hasil Perbandingan Dokumen', level=1)
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(df.columns):
            row_cells[i].text = str(row[col_name])
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def create_recommendation_highlight_docx(file_bytes, recommendations):
    doc = docx.Document(io.BytesIO(file_bytes))
    misplaced_paragraphs = [rec.get("Paragraf yang Perlu Dipindah") for rec in recommendations]
    for para in doc.paragraphs:
        if para.text.strip() in [p.strip() for p in misplaced_paragraphs if p]:
            for run in para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

# ==============================================================================
# ENDPOINTS API (Jembatan antara Frontend dan Backend)
# ==============================================================================

@app.route('/')
def home():
    """Menyajikan halaman HTML utama."""
    return render_template('index.html')

# --- Endpoint Fitur 1: Proofreading ---
@app.route('/api/proofread/analyze', methods=['POST'])
def api_proofread_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    
    try:
        document_pages = _get_text_from_flask_file(file)
        all_errors = []
        for page in document_pages:
            found_errors_on_page = proofread_with_gemini(page['teks'])
            for error in found_errors_on_page:
                all_errors.append({
                    "Kata/Frasa Salah": error['salah'],
                    "Perbaikan Sesuai KBBI": error['benar'],
                    "Pada Kalimat": error['kalimat'],
                    "Ditemukan di Halaman": page['halaman']
                })
        return jsonify(all_errors)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def _generate_proofread_files(file, file_bytes):
    """Helper internal untuk menghindari duplikasi kode di endpoint download."""
    document_pages = _extract_text_with_pages(file_bytes, file.filename.split('.')[-1].lower())
    all_errors = []
    for page in document_pages:
        found_errors = proofread_with_gemini(page['teks'])
        for error in found_errors:
             all_errors.append({
                "Kata/Frasa Salah": error['salah'],
                "Perbaikan Sesuai KBBI": error['benar'],
             })
    
    revised_data = generate_revised_docx(file_bytes, all_errors)
    highlighted_data = generate_highlighted_docx(file_bytes, all_errors)
    return revised_data, highlighted_data, file.filename

@app.route('/api/proofread/download/revised', methods=['POST'])
def api_proofread_download_revised():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read()
    
    revised_data, _, filename = _generate_proofread_files(file, file_bytes)
    
    return send_file(
        io.BytesIO(revised_data),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"revisi_{filename}"
    )

@app.route('/api/proofread/download/highlighted', methods=['POST'])
def api_proofread_download_highlighted():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read()
    
    _, highlighted_data, filename = _generate_proofread_files(file, file_bytes)
    
    return send_file(
        io.BytesIO(highlighted_data),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"highlight_{filename}"
    )

@app.route('/api/proofread/download/zip', methods=['POST'])
def api_proofread_download_zip():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read()
    
    revised_data, highlighted_data, filename = _generate_proofread_files(file, file_bytes)
    zip_data = create_zip_archive(revised_data, highlighted_data, filename)
    
    return send_file(
        io.BytesIO(zip_data),
        mimetype='application/zip',
        as_attachment=True,
        download_name=f"hasil_proofread_{filename.split('.')[0]}.zip"
    )

# --- Endpoint Fitur 2: Perbandingan Dokumen ---
def _analyze_comparison(file1_bytes, file2_bytes):
    original_paras = extract_paragraphs(file1_bytes)
    revised_paras = extract_paragraphs(file2_bytes)
    comparison_results = []
    matcher = difflib.SequenceMatcher(None, original_paras, revised_paras)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace':
            for i in range(i1, i2):
                original_para = original_paras[i]
                revised_para = revised_paras[j1 + (i - i1)] if (j1 + (i - i1)) < j2 else ""
                if revised_para:
                    word_diff = find_word_diff(original_para, revised_para)
                    comparison_results.append({
                        "Kalimat Awal": original_para,
                        "Kalimat Revisi": revised_para,
                        "Kata yang Direvisi": word_diff,
                    })
    return comparison_results

@app.route('/api/compare/analyze', methods=['POST'])
def api_compare_analyze():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Butuh dua file"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    try:
        results = _analyze_comparison(file1.read(), file2.read())
        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/compare/download', methods=['POST'])
def api_compare_download():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Butuh dua file"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    try:
        results = _analyze_comparison(file1.read(), file2.read())
        df_comparison = pd.DataFrame(results)
        
        if df_comparison.empty:
            return jsonify({"error": "Tidak ada perbedaan untuk diunduh"}), 400
            
        docx_data = create_comparison_docx(df_comparison)
        
        return send_file(
            io.BytesIO(docx_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"perbandingan_{file1.filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- Endpoint Fitur 3: Analisis Koherensi ---
@app.route('/api/coherence/analyze', methods=['POST'])
def api_coherence_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    
    try:
        full_text = _get_full_text_from_file(file)
        issues = analyze_document_coherence(full_text)
        return jsonify(issues)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- Endpoint Fitur 4: Restrukturisasi Koherensi ---
def _analyze_restructure(file):
    full_text = _get_full_text_from_file(file)
    recommendations = get_structural_recommendations(full_text)
    processed_results = []
    for rec in recommendations:
        processed_results.append({
            "Paragraf yang Perlu Dipindah": rec.get("misplaced_paragraph"),
            "Lokasi Asli": rec.get("original_section"),
            "Saran Lokasi Baru": rec.get("recommended_section")
        })
    return processed_results

@app.route('/api/restructure/analyze', methods=['POST'])
def api_restructure_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    
    try:
        results = _analyze_restructure(file)
        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/restructure/download', methods=['POST'])
def api_restructure_download():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() # Baca sekali
    
    try:
        full_text = "\n".join([p['teks'] for p in _extract_text_with_pages(file_bytes, file.filename.split('.')[-1].lower())])
        recommendations = get_structural_recommendations(full_text)
        processed_results = []
        for rec in recommendations:
            processed_results.append({
                "Paragraf yang Perlu Dipindah": rec.get("misplaced_paragraph")
            })

        if not processed_results:
             return jsonify({"error": "Tidak ada rekomendasi untuk diunduh"}), 400

        highlighted_data = create_recommendation_highlight_docx(file_bytes, processed_results)
        
        return send_file(
            io.BytesIO(highlighted_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"highlight_rekomendasi_{file.filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
if __name__ == '__main__':
    app.run(debug=True, port=5000)